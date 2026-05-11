# utils/file_handler.py

import fitz  # pymupdf
from docx import Document
from models.schemas import CVData

def extract_text_from_pdf(file_path: str) -> str:
    text_blocks = []
    # استخدام with بيقفل الملف تلقائياً بعد ما يخلص أو لو حصل error
    with fitz.open(file_path) as doc:
        for page in doc:
            text_blocks.append(page.get_text())
            
    return "\n".join(text_blocks).strip()

def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    text_blocks = []
    
    # 1. استخراج النص العادي
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_blocks.append(paragraph.text.strip())
            
    # 2. استخراج النص من الجداول (خطوة حاسمة للـ CVs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_blocks.append(cell.text.strip())
                    
    return "\n".join(text_blocks).strip()

def parse_cv_file(file_path: str, file_name: str) -> CVData:
    file_type = file_name.split(".")[-1].lower()
    
    if file_type == "pdf":
        raw_text = extract_text_from_pdf(file_path)
    elif file_type == "docx": # شلنا الـ doc لتجنب الـ Errors
        raw_text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"نوع الملف غير مدعوم: {file_type} .. برجاء رفع pdf أو docx")
    
    return CVData(
        raw_text=raw_text,
        file_name=file_name,
        file_type=file_type
    )